# -*- coding: utf-8 -*-

#pip install openai==0.28
#pip install yfinance
#pip install python-docx
import os
from dotenv import load_dotenv
import openai   #Meter IA dentro de nuestros programas
import yfinance as yf #Descargar los datos financieros de alguna compañia
import pandas as pd #Manejar datos
import matplotlib.pyplot as plt  #Hacer gráficas
from docx import Document  #Crear documentos de Word
from docx.shared import Inches #Meter gráficas en Word

"""### Configurar la API de OpenIA"""

openai.api_key = "API-KEY"

def obtener_ticker_empresa(modelo, instrucciones_sistema, nombre_empresa):
    try:
        respuesta = openai.ChatCompletion.create(
            model = modelo,
            messages=[
                {"role": "system", "content": instrucciones_sistema},
                {"role": "user", "content": f"¿Cuál es el símbolo bursátil (ticker) de la empresa {nombre_empresa} que git cotiza en bolsa en el mercado estadounidense?"}
            ],
            max_tokens=5,
            temperature=0,
            n=1,
            stop=None,
        )
        ticker = respuesta['choices'][0]['message']['content'].strip().upper()
        return ticker
    except openai.error.OpenAIError as e:
        print(f"Ocurrió un error al llamar a la API de OpenAI: {e}")

def obtener_resumen_empresa(modelo, instrucciones_sistema, nombre_empresa):
    try:
        respuesta = openai.ChatCompletion.create(
            model=modelo,
            messages=[
                {"role": "system", "content": instrucciones_sistema},
                {"role": "user", "content": f"Proporciona un resumen detallado sobre la empresa {nombre_empresa}, incluyendo sus fortalezas y debilidades actuales en el mercado."}
            ],
            max_tokens=600,
            temperature=0.7,
            n=1,
            stop=None,
        )
        resumen = respuesta['choices'][0]['message']['content'].strip()
        return resumen
    except openai.error.OpenAIError as e:
        print(f"Ocurrió un error al llamar a la API de OpenAI: {e}")
        return None

def recomendacion(modelo, instrucciones_sistema, nombre_empresa):
    try:
        respuesta = openai.ChatCompletion.create(
            model=modelo,
            messages=[
                {"role": "system", "content": "Eres un analista financiero que se encarga de analizar empresas y decidir si comprar o vender"},
                {"role": "user", "content": f"Haz una recomendación de compra o venta según lo que sabes de {nombre_empresa}"}
            ],
            max_tokens=600,
            temperature=0.7,
            n=1,
            stop=None,
        )
        resumen = respuesta['choices'][0]['message']['content'].strip()
        return resumen
    except openai.error.OpenAIError as e:
        print(f"Ocurrió un error al llamar a la API de OpenAI: {e}")
        return None

"""### Obtener datos financieros utilizando yahoo finance"""

# Función para descargar datos históricos
def descargar_datos_históricos(ticker):
    data = yf.download(ticker, period="3mo")
    if data.empty:
        print("No se pudieron descargar los datos financieros.")
        return None
    return data

data = yf.download("NVDA",period="3mo")


plt.plot(data["Close"])

def generar_grafico_medias_moviles(ticker):

  data = descargar_datos_históricos(ticker)

  data['MA15'] = data['Close'].rolling(window=15).mean()
  data["MA30"] = data["Close"].rolling(window=30).mean()


  plt.figure(figsize=(12,6))
  plt.plot(data['Close'], label='Precio de Cierre')
  plt.plot(data['MA15'], label='Media Móvil 15 días')
  plt.plot(data["MA30"], label = "Media Móvil de 30 días")
  plt.title(f'Precio de GOOGL con Medias Móviles')
  plt.xlabel('Fecha')
  plt.ylabel('Precio')
  plt.legend()
  plt.grid(True)

  # Guardar gráfico
  nombre_grafico = f"grafico_{ticker}.png"
  plt.savefig(nombre_grafico)
  plt.close()
  return nombre_grafico

generar_grafico_medias_moviles("PEP")

"""### Automatizar nuestro reporte utilizando la libreria de word"""

def crear_documento_word(nombre_empresa, resumen, nombre_grafico,recomendacion):
    doc = Document()
    doc.add_heading(f"Análisis de {nombre_empresa}", 0)

    doc.add_heading("Resumen de la Empresa", level=1)
    doc.add_paragraph(resumen)

    doc.add_heading("Gráfico de Precio con Medias Móviles", level=1)
    doc.add_picture(nombre_grafico, width=Inches(6))

    #Recomendacion
    doc.add_heading("Recomendación de compra y venta", level=1)
    doc.add_paragraph(recomendacion)



    nombre_documento = f"Analisis_{nombre_empresa.replace(' ', '_')}.docx"
    doc.save(nombre_documento)
    return nombre_documento

def main():
  nombre_empresa = input("Ingrese el nombre de la empresa: ")

  instrucciones_sistema_ticker =  '''
    Eres un experto en finanzas y mercados bursátiles. Tu tarea es proporcionar el símbolo bursátil (ticker) oficial de una empresa que cotiza en bolsa en el mercado estadounidense. No des una respuesta larga
    la única palabra que debes responder es el ticker. Por ejemplo si te dicen "Tesla" debes responder "TSLA".
    '''


  instrucciones_sistema_resumen = '''
    Eres un analista financiero experto. Proporciona un resumen detallado sobre la empresa indicada, incluyendo sus fortalezas y debilidades actuales en el mercado.
    '''

  modelo = "gpt-3.5-turbo"

  ticker = obtener_ticker_empresa(modelo, instrucciones_sistema_ticker, nombre_empresa)
  if ticker is None:
    print("No se pudo obtener el ticker de la empresa.")
    return

  #Generar gráfico
  nombre_grafico = generar_grafico_medias_moviles(ticker)

  #Obtener resumen de la empresa

  resumen = obtener_resumen_empresa(modelo, instrucciones_sistema_resumen, nombre_empresa)
  if resumen is None:
    print("No se pudo obtener el resumen de la empresa.")
    return

  #Obtener recomendacion

  recomendacion_texto = recomendacion(modelo, instrucciones_sistema_resumen, nombre_empresa)

  #Crear documento Word

  nombre_documento = crear_documento_word(nombre_empresa, resumen, nombre_grafico, recomendacion_texto)
  print(f"El documento Word ha sido creado: {nombre_documento}")

if __name__ == "__main__":
    main()